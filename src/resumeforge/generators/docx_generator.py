"""DOCX generator for resume output."""

from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

if TYPE_CHECKING:
    from resumeforge.schemas.blackboard import Blackboard


class DocxGenerator:
    """Generates DOCX output from blackboard state."""
    
    def __init__(self, template_path: str | None = None):
        """
        Initialize generator.
        
        Args:
            template_path: Optional path to DOCX template
        """
        self.template_path = template_path
    
    def generate(self, blackboard: "Blackboard", output_path: Path) -> None:
        """
        Generate DOCX file from blackboard.
        
        Args:
            blackboard: Blackboard with resume draft
            output_path: Path to output DOCX file
        """
        if not blackboard.resume_draft:
            raise ValueError("Cannot generate DOCX: resume_draft is None")
        
        # Initialize document
        if self.template_path and Path(self.template_path).exists():
            doc = Document(self.template_path)
        else:
            doc = Document()
            self._setup_default_styles(doc)
        
        # Add sections from resume draft
        for section in blackboard.resume_draft.sections:
            self._add_section(doc, section)
        
        # Save document
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
    
    def _setup_default_styles(self, doc: Document) -> None:
        """
        Configure default document styles.
        
        Args:
            doc: Document to configure
        """
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Arial"
        font.size = Pt(11)
    
    def _add_section(self, doc: Document, section) -> None:
        """
        Add a section to the document.
        
        Args:
            doc: Document to add section to
            section: ResumeSection to add
        """
        # Add heading
        doc.add_heading(section.name, level=1)
        
        # Parse markdown content and add paragraphs
        lines = section.content.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if line.startswith("- ") or line.startswith("* "):
                # Bullet point - remove bullet marker
                bullet_text = line[2:].strip()
                if bullet_text:
                    p = doc.add_paragraph(bullet_text, style="List Bullet")
            elif line.startswith("#"):
                # Markdown heading - convert to Word heading
                level = len(line) - len(line.lstrip("#"))
                heading_text = line.lstrip("#").strip()
                if heading_text:
                    doc.add_heading(heading_text, level=min(level, 3))
            else:
                # Regular paragraph
                doc.add_paragraph(line)
